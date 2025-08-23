import os
import pandas as pd
from reportlab.pdfgen import canvas
from docx import Document

# Ensure the input_files folder exists
os.makedirs("input_files", exist_ok=True)

# Common data
data = {
    "Name": ["Ravi", "Priya", "Amit"],
    "Age": [25, 28, 30],
    "City": ["Delhi", "Mumbai", "Bangalore"]
}

# 1. CSV
df = pd.DataFrame(data)
csv_path = "input_files/sample.csv"
df.to_csv(csv_path, index=False)
print(f"[CREATED] {csv_path}")

# 2. Excel (.xlsx)
xlsx_path = "input_files/sample.xlsx"
df.to_excel(xlsx_path, index=False)
print(f"[CREATED] {xlsx_path}")

# 3. Excel (.xls)
# xls_path = "input_files/sample.xls"
# df.to_excel(xls_path, index=False)
# print(f"[CREATED] {xls_path}")

# 4. PDF
pdf_path = "input_files/sample.pdf"
c = canvas.Canvas(pdf_path)
c.drawString(100, 750, "Name: Ravi, Age: 25, City: Delhi")
c.drawString(100, 730, "Name: Priya, Age: 28, City: Mumbai")
c.drawString(100, 710, "Name: Amit, Age: 30, City: Bangalore")
c.save()
print(f"[CREATED] {pdf_path}")

# 5. DOCX
docx_path = "input_files/sample.docx"
doc = Document()
doc.add_heading('Sample People Data', level=1)
doc.add_paragraph('Ravi, 25, Delhi')
doc.add_paragraph('Priya, 28, Mumbai')
doc.add_paragraph('Amit, 30, Bangalore')
doc.save(docx_path)
print(f"[CREATED] {docx_path}")

print("\n✅ All sample files generated in 'input_files/' folder.")
